from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

import matplotlib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Emu, Inches, Pt, RGBColor, Twips
from reportlab.lib.utils import ImageReader

from app.db.models import Exam, ExamVersion, ExamVersionItem, Item

matplotlib.use("Agg")
import matplotlib.pyplot as plt

# Tamaño carta: 8.5" x 11"
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 0.5
GUTTER_IN = 0.3
CONTENT_WIDTH_IN = PAGE_WIDTH_IN - 2 * MARGIN_IN
COLUMN_WIDTH_IN = (CONTENT_WIDTH_IN - GUTTER_IN) / 2  # ~3.6"

# Altura inline para ecuaciones y valores numéricos (puntos)
INLINE_IMG_HEIGHT_PT = 18.0

# Tamaño de fuente para renderizar LaTeX a imagen (matplotlib)
LATEX_FONT_PT = 16.0

BACKEND_DIR = Path(__file__).resolve().parents[3]
ASSETS_DIR = BACKEND_DIR / "data" / "input" / "item_assets"
ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def _parse_storage_doc(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return {"type": "doc", "content": []}
        try:
            parsed = json.loads(text)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass
        return {
            "type": "doc",
            "content": [{"type": "paragraph", "content": [{"type": "text", "text": value}]}],
        }
    return {"type": "doc", "content": []}


def _asset_path_from_src(src: str) -> Path | None:
    src = src.strip()
    if not src:
        return None
    assets_root = BACKEND_DIR / "data" / "input"

    if src.startswith("/assets/"):
        rel = src.removeprefix("/assets/").strip("/")
        path = assets_root / rel
        return path if path.exists() else None

    marker = "/assets/"
    idx = src.find(marker)
    if idx >= 0:
        rel = src[idx + len(marker):].strip("/")
        path = assets_root / rel
        return path if path.exists() else None
    return None


def _render_latex_to_png(latex: str) -> Path | None:
    clean = str(latex or "").strip()
    if not clean:
        return None
    while clean.startswith("$") and clean.endswith("$") and len(clean) >= 2:
        clean = clean[1:-1].strip()
    if not clean:
        return None

    output_path = ASSETS_DIR / f"eq_{uuid4().hex}.png"
    try:
        fig = plt.figure(figsize=(0.01, 0.01))
        fig.patch.set_alpha(0)
        text = fig.text(0, 0, f"${clean}$", fontsize=LATEX_FONT_PT)
        fig.canvas.draw()
        bbox = text.get_window_extent(renderer=fig.canvas.get_renderer()).expanded(1.08, 1.25)
        dpi = 160
        width_in = max(bbox.width / dpi, 0.2)
        height_in = max(bbox.height / dpi, 0.12)
        fig.set_size_inches(width_in, height_in)
        text.set_position((0.02, 0.04))
        fig.savefig(output_path, dpi=dpi, transparent=True, format="png",
                    bbox_inches="tight", pad_inches=0.01)
        plt.close(fig)
        return output_path if output_path.exists() else None
    except Exception:  # noqa: BLE001
        plt.close("all")
        return None


def _image_pixel_size(path: Path) -> tuple[int, int]:
    try:
        reader = ImageReader(str(path))
        iw, ih = reader.getSize()
        return int(iw), int(ih)
    except Exception:  # noqa: BLE001
        return 0, 0


def _merge_inline_image_paragraphs(nodes: list[Any]) -> list[Any]:
    """
    Fusiona párrafos de imagen pequeña (alto < 160px) con el párrafo de texto anterior
    para que se rendericen inline en lugar de como bloque separado.
    """
    result: list[Any] = []
    for node in nodes:
        if not isinstance(node, dict) or node.get("type") != "paragraph":
            result.append(node)
            continue

        content = node.get("content") or []
        # ¿Es un párrafo de imagen pura?
        if len(content) != 1 or not isinstance(content[0], dict) or content[0].get("type") != "image":
            result.append(node)
            continue

        src = (content[0].get("attrs") or {}).get("src", "")
        path = _asset_path_from_src(str(src)) if isinstance(src, str) else None
        if not path:
            result.append(node)
            continue

        _iw, ih = _image_pixel_size(path)
        is_small = 0 < ih < 160

        if is_small and result and isinstance(result[-1], dict) and result[-1].get("type") == "paragraph":
            merged_content = list(result[-1].get("content") or []) + list(content)
            result[-1] = {**result[-1], "content": merged_content}
        else:
            result.append(node)

    return result


def _no_spacing(p: Any) -> Any:
    """Elimina espacio antes y después del párrafo."""
    p.paragraph_format.space_before = Twips(0)
    p.paragraph_format.space_after = Twips(0)
    return p


def _remove_image_border(run: Any) -> None:
    """Quita el recuadro/borde de una imagen inline insertada con run.add_picture()."""
    from lxml import etree

    NS_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    NS_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

    r_elem = run._r
    drawing = r_elem.find(qn("w:drawing"))
    if drawing is None:
        return

    inline = drawing.find(f"{{{NS_WP}}}inline")
    if inline is None:
        return

    graphic = inline.find(f"{{{NS_A}}}graphic")
    if graphic is None:
        return

    graphic_data = graphic.find(f"{{{NS_A}}}graphicData")
    if graphic_data is None:
        return

    pic_el = graphic_data.find(f"{{{NS_PIC}}}pic")
    if pic_el is None:
        return

    sp_pr = pic_el.find(f"{{{NS_PIC}}}spPr")
    if sp_pr is None:
        return

    # Quitar bordes existentes y agregar línea sin relleno
    for ln in sp_pr.findall(f"{{{NS_A}}}ln"):
        sp_pr.remove(ln)
    ln_el = etree.SubElement(sp_pr, f"{{{NS_A}}}ln")
    etree.SubElement(ln_el, f"{{{NS_A}}}noFill")

    # Agregar solidFill transparente para evitar fondo de placeholder
    for solidFill in sp_pr.findall(f"{{{NS_A}}}solidFill"):
        sp_pr.remove(solidFill)
    solidFill_el = etree.SubElement(sp_pr, f"{{{NS_A}}}solidFill")
    srgbClr = etree.SubElement(solidFill_el, f"{{{NS_A}}}srgbClr")
    srgbClr.set("val", "FFFFFF")
    alpha_el = etree.SubElement(srgbClr, f"{{{NS_A}}}alpha")
    alpha_el.set("val", "0")

    # Deshabilitar el marco de recorte en cNvPicPr
    cNvPicPr = pic_el.find(f"{{{NS_PIC}}}nvPicPr/{{{NS_PIC}}}cNvPicPr")
    if cNvPicPr is not None:
        cNvPicPr.set("preferRelativeResize", "0")


def _add_picture_no_border(run: Any, path: Path, **kwargs: Any) -> None:
    """Inserta una imagen en el run y le quita el borde."""
    run.add_picture(str(path), **kwargs)
    _remove_image_border(run)


def _set_page_size_letter(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Emu(int(PAGE_WIDTH_IN * 914400))
    section.page_height = Emu(int(PAGE_HEIGHT_IN * 914400))
    section.left_margin = Emu(int(MARGIN_IN * 914400))
    section.right_margin = Emu(int(MARGIN_IN * 914400))
    section.top_margin = Emu(int(MARGIN_IN * 914400))
    section.bottom_margin = Emu(int(MARGIN_IN * 914400))

    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:cols")):
        sectPr.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(int(0.3 * 1440)))
    cols.set(qn("w:sep"), "1")
    sectPr.append(cols)


def _set_paragraph_border(p: Any, border_kind: str, *, color: str = "B8C0CC", size: int = 6) -> None:
    """Agrega un borde simple a un párrafo Word."""
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    existing = pBdr.find(qn(f"w:{border_kind}"))
    if existing is not None:
        pBdr.remove(existing)

    border = OxmlElement(f"w:{border_kind}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)
    pBdr.append(border)


def _add_question_separator(
    doc: Document,
    *,
    border_kind: str,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    """Inserta un párrafo vacío usado como separador visual de cada pregunta."""
    p = _no_spacing(doc.add_paragraph())
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    _set_paragraph_border(p, border_kind)


def _fill_cell_content(p: Any, cell_node: dict[str, Any], is_header: bool, cell_height_pt: float = 9.0) -> None:
    """Rellena un párrafo de celda con contenido mixto: texto y mathInline renderizado."""
    def walk(node: Any) -> None:
        if not isinstance(node, dict):
            return
        ntype = node.get("type")
        if ntype == "text":
            txt = node.get("text") or ""
            if txt:
                r = p.add_run(str(txt))
                r.font.size = Pt(8)
                if is_header:
                    r.bold = True
        elif ntype == "hardBreak":
            p.add_run().add_break()
        elif ntype == "mathInline":
            latex = (node.get("attrs") or {}).get("latex", "")
            if latex:
                path = _render_latex_to_png(str(latex))
                if path:
                    r = p.add_run()
                    _add_picture_no_border(r, path, height=Pt(cell_height_pt))
        else:
            for child in (node.get("content") or []):
                walk(child)
    walk(cell_node)


def _add_table_node(doc: Document, node: dict[str, Any], column_width_in: float = COLUMN_WIDTH_IN) -> None:
    rows_nodes = [r for r in (node.get("content") or []) if isinstance(r, dict) and r.get("type") == "tableRow"]
    if not rows_nodes:
        return

    # Contar columnas
    num_cols = max(
        len([c for c in (r.get("content") or []) if isinstance(c, dict)])
        for r in rows_nodes
    )
    if num_cols == 0:
        return

    safe_table_width_in = max(column_width_in - 0.12, 1.0)
    col_width_twips = max(int((safe_table_width_in / num_cols) * 1440), 1)
    col_width_pct = max(int(5000 / num_cols), 1)

    table = doc.add_table(rows=len(rows_nodes), cols=num_cols)
    table.style = "Table Grid"
    table.autofit = True

    # Ancho total relativo al contenedor (la columna Word), para que Google Docs
    # y Word puedan envolver el texto sin fijar un ancho absoluto que desborde.
    tbl = table._tbl
    tblPr = tbl.tblPr
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "autofit")
    tblPr.append(tblLayout)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    tblGrid = tbl.tblGrid
    if tblGrid is not None:
        for grid_col in tblGrid.findall(qn("w:gridCol")):
            grid_col.set(qn("w:w"), str(col_width_twips))

    for r_idx, row_node in enumerate(rows_nodes):
        cells_nodes = [c for c in (row_node.get("content") or []) if isinstance(c, dict)]
        is_header_row = any(c.get("type") == "tableHeader" for c in cells_nodes)

        for c_idx in range(num_cols):
            cell = table.cell(r_idx, c_idx)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:tcW")):
                tcPr.remove(existing)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_width_pct))
            tcW.set(qn("w:type"), "pct")
            tcPr.append(tcW)

            # Limpiar contenido previo y rellenar
            p = _no_spacing(cell.paragraphs[0])
            for run in list(p.runs):
                run._r.getparent().remove(run._r)

            if c_idx < len(cells_nodes):
                _fill_cell_content(p, cells_nodes[c_idx], is_header=is_header_row)


def _collect_plain_text(node: Any, parts: list[str]) -> None:
    if not isinstance(node, dict):
        return
    node_type = node.get("type")
    if node_type == "text":
        text = node.get("text")
        if text:
            parts.append(str(text))
        return
    if node_type == "hardBreak":
        parts.append(" ")
        return
    if node_type == "mathInline":
        latex = (node.get("attrs") or {}).get("latex", "")
        if latex:
            # Quitar delimitadores $ para mostrar el latex limpio en la celda
            clean = str(latex).strip()
            while clean.startswith("$") and clean.endswith("$") and len(clean) >= 2:
                clean = clean[1:-1].strip()
            parts.append(clean)
        return
    for child in (node.get("content") or []):
        _collect_plain_text(child, parts)


def _render_paragraph_node(
    doc: Document,
    node: dict[str, Any],
    body_size_pt: float,
    column_width_in: float,
    prefix: str = "",
) -> None:
    """
    Renderiza un nodo párrafo Tiptap como un único párrafo Word con runs múltiples.
    Texto, mathInline e imágenes pequeñas van en el mismo párrafo (inline).
    Imágenes grandes van en un párrafo propio.
    """
    content = node.get("content") or []

    # Párrafo de imagen pura (grande): va en su propio párrafo
    if len(content) == 1 and isinstance(content[0], dict) and content[0].get("type") == "image":
        src = (content[0].get("attrs") or {}).get("src", "")
        path = _asset_path_from_src(str(src)) if isinstance(src, str) else None
        if path:
            p = _no_spacing(doc.add_paragraph())
            run = p.add_run()
            _add_picture_no_border(run, path, width=Inches(min(column_width_in - 0.1, 3.5)))
        return

    # Párrafo mixto: crear UN párrafo y agregar runs secuencialmente
    p = _no_spacing(doc.add_paragraph())

    if prefix:
        r = p.add_run(prefix)
        r.bold = True
        r.font.size = Pt(body_size_pt)

    for child in content:
        if not isinstance(child, dict):
            continue
        child_type = child.get("type")

        if child_type == "text":
            txt = child.get("text") or ""
            if not txt:
                continue
            r = p.add_run(str(txt))
            r.font.size = Pt(body_size_pt)
            for mark in (child.get("marks") or []):
                if not isinstance(mark, dict):
                    continue
                mt = mark.get("type")
                if mt == "bold":
                    r.bold = True
                elif mt == "italic":
                    r.italic = True
                elif mt == "underline":
                    r.underline = True

        elif child_type == "hardBreak":
            r = p.add_run()
            r.add_break()

        elif child_type == "mathInline":
            latex = (child.get("attrs") or {}).get("latex", "")
            if latex:
                path = _render_latex_to_png(str(latex))
                if path:
                    r = p.add_run()
                    _add_picture_no_border(r, path, height=Pt(INLINE_IMG_HEIGHT_PT))

        elif child_type == "image":
            src = (child.get("attrs") or {}).get("src", "")
            path = _asset_path_from_src(str(src)) if isinstance(src, str) else None
            if path:
                _iw, ih = _image_pixel_size(path)
                if ih > 0 and ih < 160:
                    # Imagen pequeña inline (número, valor)
                    r = p.add_run()
                    _add_picture_no_border(r, path, height=Pt(INLINE_IMG_HEIGHT_PT))
                else:
                    # Imagen grande
                    r = p.add_run()
                    _add_picture_no_border(r, path, width=Inches(min(column_width_in - 0.1, 3.5)))


def _render_tiptap_to_docx(
    value: Any,
    doc: Document,
    body_size_pt: float,
    column_width_in: float,
    prefix: str = "",
) -> None:
    """Renderiza un doc Tiptap completo al documento Word."""
    tiptap_doc = _parse_storage_doc(value)
    raw_nodes = tiptap_doc.get("content") or []
    nodes = _merge_inline_image_paragraphs(raw_nodes)

    first = True
    for node in nodes:
        if not isinstance(node, dict):
            continue
        node_type = node.get("type")

        if node_type == "image":
            # Imagen a nivel de doc (no dentro de paragraph): siempre bloque
            src = (node.get("attrs") or {}).get("src", "")
            path = _asset_path_from_src(str(src)) if isinstance(src, str) else None
            if path:
                p = _no_spacing(doc.add_paragraph())
                if prefix and first:
                    r = p.add_run(prefix)
                    r.bold = True
                    r.font.size = Pt(body_size_pt)
                r = p.add_run()
                _add_picture_no_border(r, path, width=Inches(min(column_width_in - 0.1, 3.5)))
                first = False

        elif node_type == "table":
            _add_table_node(doc, node, column_width_in=column_width_in)

        elif node_type == "paragraph":
            _render_paragraph_node(
                doc, node, body_size_pt, column_width_in,
                prefix=prefix if first else "",
            )
            first = False

        else:
            # Otros nodos (bulletList, etc.): extraer texto plano
            parts: list[str] = []
            _collect_plain_text(node, parts)
            text = " ".join(p for p in parts if p).strip()
            if text:
                p = _no_spacing(doc.add_paragraph())
                r = p.add_run((prefix if first else "") + text)
                r.font.size = Pt(body_size_pt)
                first = False


def _mapped_options_for_version(item: Item, option_map: dict[str, str] | None) -> dict[str, Any]:
    options = item.options or {}
    if not option_map:
        return {k: options.get(k, "") for k in ("A", "B", "C", "D")}

    mapped: dict[str, Any] = {}
    for original_label, mapped_label in option_map.items():
        if mapped_label in {"A", "B", "C", "D"}:
            mapped[mapped_label] = options.get(original_label, "")
    for label in ("A", "B", "C", "D"):
        mapped.setdefault(label, options.get(label, ""))
    return mapped


def build_exam_version_docx(
    *,
    exam: Exam,
    version: ExamVersion,
    version_items: list[ExamVersionItem],
    items_by_id: dict[int, Item],
) -> bytes:
    doc = Document()
    _set_page_size_letter(doc)

    title_p = doc.add_heading(f"Cuadernillo: {exam.title}", level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Código de examen: {version.exam_code}   |   Versión: {version.version_code}").bold = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for version_row in version_items:
        item = items_by_id.get(version_row.item_id)
        if item is None:
            continue

        q_heading = _no_spacing(doc.add_paragraph())
        _set_paragraph_border(q_heading, "top")
        q_run = q_heading.add_run(f"Pregunta {version_row.question_number}")
        q_run.bold = True
        q_run.font.size = Pt(10)
        q_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        _render_tiptap_to_docx(item.statement, doc, body_size_pt=9, column_width_in=COLUMN_WIDTH_IN)

        mapped_options = _mapped_options_for_version(item, version_row.option_map_json)
        for label in ("A", "B", "C", "D"):
            _render_tiptap_to_docx(
                mapped_options.get(label, ""), doc,
                body_size_pt=9, column_width_in=COLUMN_WIDTH_IN,
                prefix=f"{label}. ",
            )

        _add_question_separator(doc, border_kind="bottom")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
